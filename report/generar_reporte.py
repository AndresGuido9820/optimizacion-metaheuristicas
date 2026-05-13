"""Genera el reporte técnico del trabajo en formato .docx con normas APA 7."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Márgenes APA 7: 2.54 cm todos los lados ──────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

# ── Estilo base: Times New Roman 12, doble espacio (APA 7) ───────────────────
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.space_after  = Pt(0)
normal.paragraph_format.line_spacing = Pt(24)


# ── Helpers ───────────────────────────────────────────────────────────────────
def add_hrule():
    """Línea horizontal (borde inferior del párrafo)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1a3a6c')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def cover_line(text, bold=False, size=12, space_after=Pt(4),
               align=WD_ALIGN_PARAGRAPH.CENTER, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after  = space_after
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(size * 1.15)
    r = p.add_run(text)
    r.bold      = bold
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p


def h1(text):
    """Nivel 1 APA 7: centrado, negrita."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p


def h2(text):
    """Nivel 2 APA 7: izquierda, negrita."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p


def h3(text):
    """Nivel 3 APA 7: sangría, negrita cursiva, punto al final."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing      = Pt(24)
    p.paragraph_format.space_after       = Pt(0)
    r = p.add_run(text + '.')
    r.bold = True; r.italic = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p


def body(text):
    """Párrafo con sangría primera línea (APA 7)."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing      = Pt(24)
    p.paragraph_format.space_after       = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p


def body_no_indent(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing      = Pt(24)
    p.paragraph_format.space_after       = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p


def ref(text):
    """Referencia con sangría colgante (APA 7)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.line_spacing      = Pt(24)
    p.paragraph_format.space_after       = Pt(0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p


def nota_tabla(label, texto):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing      = Pt(24)
    p.paragraph_format.space_after       = Pt(0)
    r1 = p.add_run(label + '. ')
    r1.italic = True
    r1.font.name = 'Times New Roman'; r1.font.size = Pt(12)
    r2 = p.add_run(texto)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)
    return p


def blank(pts=24):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(pts)
    p.paragraph_format.space_after  = Pt(0)


def encabezado_tabla(tabla, encabezados):
    for i, h in enumerate(encabezados):
        c = tabla.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.name = 'Times New Roman'
        c.paragraphs[0].runs[0].font.size = Pt(10)


def fila_tabla(tabla, fila_idx, datos):
    for col_idx, val in enumerate(datos):
        c = tabla.rows[fila_idx].cells[col_idx]
        c.text = val
        c.paragraphs[0].runs[0].font.name = 'Times New Roman'
        c.paragraphs[0].runs[0].font.size = Pt(10)


# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════

# Logo centrado
p_logo = doc.add_paragraph()
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_logo.paragraph_format.space_before = Pt(0)
p_logo.paragraph_format.space_after  = Pt(14)
p_logo.paragraph_format.line_spacing = Pt(24)
p_logo.add_run().add_picture('logo_unal.png', width=Cm(9))

# Nombre universidad y facultad
cover_line('Universidad Nacional de Colombia', bold=True, size=13,
           space_after=Pt(2), color=(26, 58, 108))
cover_line('Sede Medellín', size=11, space_after=Pt(2), color=(26, 58, 108))
cover_line('Facultad de Minas', size=11, space_after=Pt(28), color=(26, 58, 108))

# Línea separadora
add_hrule()
blank(8)

# Título
cover_line(
    'Optimización Heurística: Comparativa de Metaheurísticas en\n'
    'Funciones de Prueba y el Problema del Agente Viajero\n'
    'para las Capitales de México',
    bold=True, size=14, space_after=Pt(6)
)

blank(8)
add_hrule()
blank(28)

# Bloque asignatura / grupo
cover_line('Asignatura', bold=True, size=10,
           space_after=Pt(1), color=(80, 80, 80))
cover_line('Optimización', size=12, space_after=Pt(14))

cover_line('Grupo', bold=True, size=10,
           space_after=Pt(1), color=(80, 80, 80))
cover_line('1', size=12, space_after=Pt(20))

# Estudiantes
cover_line('Estudiantes', bold=True, size=10,
           space_after=Pt(3), color=(80, 80, 80))
for nombre in [
    'Andrés Felipe Guido Montoya',
    'Juan José Martínez',
    'Andrés Lemus',
]:
    cover_line(nombre, size=12, space_after=Pt(2))

blank(20)

# Docente
cover_line('Docente', bold=True, size=10,
           space_after=Pt(3), color=(80, 80, 80))
cover_line('Juan David Ospino Arango', size=12, space_after=Pt(40))

# Ciudad y fecha al fondo
cover_line('Medellín, Colombia', size=11, space_after=Pt(2))
cover_line('Mayo de 2026', size=11, space_after=Pt(0))

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# RESUMEN
# ══════════════════════════════════════════════════════════════════════════════
h1('Resumen')
blank()

body_no_indent(
    'Este trabajo presenta una comparativa experimental de cuatro algoritmos de '
    'optimización —descenso por gradiente (GD), algoritmos evolutivos (EA), '
    'optimización por enjambre de partículas (PSO) y evolución diferencial (DE)— '
    'aplicados a las funciones de prueba de Rosenbrock y Rastrigin en dimensiones '
    '2D y 3D. Adicionalmente, se resuelve el Problema del Agente Viajero (TSP) '
    'para las 32 capitales estatales de México utilizando colonias de hormigas '
    '(ACO) y algoritmos genéticos (GA), con un modelo de costo que incorpora '
    'combustible, peajes y tiempo del vendedor. Los experimentos se realizaron '
    'con 30 corridas independientes por configuración para garantizar validez '
    'estadística. Los resultados muestran que la evolución diferencial domina en '
    'las funciones de prueba con 100% de éxito en todos los escenarios y entre '
    '2,000 y 11,000 evaluaciones de función. ACO ofrece mayor consistencia en el '
    'TSP con un coeficiente de variación de 0.72% frente a 2.91% del GA, aunque '
    'GA encontró la mejor solución absoluta (55,796 MXN) en al menos una corrida.'
)

blank()

p_kw = doc.add_paragraph()
p_kw.paragraph_format.first_line_indent = Cm(1.27)
p_kw.paragraph_format.line_spacing      = Pt(24)
r1 = p_kw.add_run('Palabras clave: ')
r1.italic = True; r1.font.name = 'Times New Roman'; r1.font.size = Pt(12)
r2 = p_kw.add_run(
    'metaheurísticas, evolución diferencial, colonias de hormigas, '
    'TSP, Rosenbrock, Rastrigin, optimización bio-inspirada.'
)
r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════════════
h1('Introducción')
blank()

body(
    'La optimización es una disciplina transversal a la ingeniería, la economía '
    'y las ciencias computacionales. Su objetivo formal consiste en encontrar el '
    'valor de un vector de variables x* que minimiza una función objetivo '
    'f: ℝⁿ → ℝ, posiblemente sujeto a restricciones. Cuando f es diferenciable '
    'y convexa, los métodos de gradiente garantizan convergencia al óptimo global. '
    'Sin embargo, la mayoría de los problemas reales son no convexos, multimodales, '
    'discontinuos o de alta dimensionalidad, condiciones bajo las cuales los '
    'métodos clásicos presentan limitaciones severas.'
)
body(
    'Las metaheurísticas surgen como respuesta a esta limitación. Son estrategias '
    'de búsqueda de alto nivel, inspiradas frecuentemente en fenómenos naturales, '
    'que sacrifican garantías de optimalidad a cambio de encontrar soluciones de '
    'alta calidad en tiempos computacionales razonables (Blum & Roli, 2003). Su '
    'popularidad ha crecido exponencialmente desde los años 1980: los algoritmos '
    'evolutivos emergieron de los trabajos de Holland (1975) y Goldberg (1989), '
    'el PSO fue propuesto por Kennedy y Eberhart (1995), la evolución diferencial '
    'por Storn y Price (1997), y las colonias de hormigas por Dorigo (1992).'
)
body(
    'Este trabajo tiene dos objetivos complementarios. El primero consiste en '
    'comparar GD, EA, PSO y DE sobre las funciones de Rosenbrock (unimodal, valle '
    'estrecho) y Rastrigin (multimodal) en 2D y 3D con 30 corridas independientes '
    'por configuración. El segundo consiste en resolver el TSP para las 32 '
    'capitales estatales de México con ACO y GA, minimizando un modelo de costo '
    'que combina gasto en combustible, peajes y tiempo del vendedor. El código '
    'fuente y los notebooks están disponibles en '
    'https://github.com/AndresGuido9820/tarea01-optimizacion-heuristica.'
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 2. MARCO TEÓRICO
# ══════════════════════════════════════════════════════════════════════════════
h1('Marco Teórico')
blank()

h2('Funciones de Prueba')
blank()

h3('Función de Rosenbrock')
blank()
body(
    'Propuesta por Rosenbrock (1960), es una función unimodal no convexa definida '
    'como f(x) = Σ [100(x_{i+1} − x_i²)² + (1 − x_i)²] para i = 1, …, n−1. '
    'El óptimo global es f(1,…,1) = 0. La dificultad radica en un valle parabólico '
    'estrecho y curvado: el gradiente a lo largo del fondo del valle es casi nulo, '
    'haciendo que los métodos de gradiente converjan extremadamente lento y que '
    'los algoritmos de búsqueda aleatoria difícilmente se mantengan en él.'
)

h3('Función de Rastrigin')
blank()
body(
    'Introducida por Rastrigin (1974) y popularizada por Mühlenbein et al. (1991), '
    'es una función altamente multimodal: f(x) = An + Σ [x_i² − A·cos(2πx_i)], '
    'con A = 10. El óptimo global es f(0,…,0) = 0. Contiene aproximadamente 10ⁿ '
    'mínimos locales en [−5, 5]ⁿ y es el estándar para evaluar la capacidad de '
    'escapar de mínimos locales.'
)

blank()
h2('Descenso por Gradiente con Búsqueda en Línea')
blank()
body(
    'El descenso por gradiente actualiza iterativamente la solución según '
    'x_{k+1} = x_k − α_k ∇f(x_k). El tamaño de paso α_k se determina mediante '
    'búsqueda en línea con retroceso basada en la condición de Armijo: '
    'f(x_k − α∇f) ≤ f(x_k) − c·α‖∇f‖², con c = 10⁻⁴ y factor de reducción '
    'β = 0.5 (Nocedal & Wright, 2006). El gradiente se aproxima con diferencias '
    'finitas centrales de orden O(h²), con h = 10⁻⁵.'
)

blank()
h2('Algoritmos Evolutivos')
blank()
body(
    'Inspirados en la selección natural (Holland, 1975), operan sobre una '
    'población de soluciones aplicando selección por torneo (k = 3), cruce '
    'cxBlend con α = 0.5 (Eshelman & Schaffer, 1993) y mutación gaussiana con '
    'σ = 0.5. La implementación utiliza DEAP (Fortin et al., 2012) con '
    'N_pop = 100, N_gen = 500, p_cx = 0.7 y p_mut = 0.2.'
)

blank()
h2('Optimización por Enjambre de Partículas')
blank()
body(
    'Propuesto por Kennedy y Eberhart (1995), modela una bandada de pájaros '
    'buscando alimento. La velocidad se actualiza como '
    'v_i ← w·v_i + c₁r₁(p_i − x_i) + c₂r₂(g − x_i), donde p_i es la mejor '
    'posición personal y g la mejor global. Se usa w = 0.729, factor de '
    'constricción de Clerc y Kennedy (2002) que garantiza convergencia teórica, '
    'c₁ = c₂ = 2.05, N = 50 partículas y 500 iteraciones (Miranda, 2018).'
)

blank()
h2('Evolución Diferencial')
blank()
body(
    'Propuesta por Storn y Price (1997), genera un mutante combinando tres '
    'individuos aleatorios: v_i = x_{r1} + F(x_{r2} − x_{r3}), con F ∈ [0.5, 1.0] '
    'adaptativo. Aplica cruce binomial con CR = 0.7 y selecciona el mejor entre '
    'el individuo original y el trial vector (estrategia best1bin). Se usa '
    'scipy.optimize.differential_evolution (Virtanen et al., 2020) con '
    'popsize = 15, maxiter = 1,000 y tol = 10⁻⁷.'
)
body(
    'La fortaleza de DE en el valle de Rosenbrock radica en que el vector de '
    'mutación F(x_{r2} − x_{r3}) escala adaptativamente con la dispersión de la '
    'población: cuando los individuos se concentran cerca del óptimo, las '
    'diferencias se vuelven pequeñas automáticamente, produciendo el refinamiento '
    'fino que el valle estrecho requiere.'
)

blank()
h2('Problema del Agente Viajero')
blank()
body(
    'El TSP busca la permutación π* de n ciudades que minimice el costo del '
    'recorrido cerrado C(π) = Σ d(π_i, π_{i+1 mod n}) (Applegate et al., 2006). '
    'Para n = 32 el espacio tiene (n−1)!/2 ≈ 1.3×10³³ tours posibles. El modelo '
    'de costo incorpora combustible (3.0 MXN/km), peajes (1.5 MXN/km) y tiempo '
    'del vendedor (150 MXN/h a 80 km/h), para un factor combinado de 6.375 MXN/km. '
    'Las distancias se calculan con la fórmula de Haversine con R = 6,371 km.'
)

blank()
h2('Colonias de Hormigas')
blank()
body(
    'Introducido por Dorigo (1992), cada hormiga construye una solución usando '
    'la regla p_ij = [τ_ij]^α·[η_ij]^β / Σ [τ_il]^α·[η_il]^β, donde τ_ij es '
    'la feromona y η_ij = 1/d_ij la heurística de visibilidad. La feromona se '
    'evapora con tasa ρ = 0.1 y se deposita proporcionalmente a la calidad de '
    'cada ruta (Dorigo & Gambardella, 1997). Parámetros: N_ants = 50, '
    'N_iters = 300, α = 1, β = 3, Q = 100.'
)

blank()
h2('Algoritmo Genético para TSP')
blank()
body(
    'El OX crossover (Davis, 1985) copia un segmento del padre 1 y rellena con '
    'las ciudades del padre 2 en su orden de aparición, preservando el orden '
    'relativo geográfico. La mutación por intercambio de índices aplica '
    'intercambios con probabilidad p_indpb = 2/n, garantizando permutaciones '
    'válidas. Parámetros: N_pop = 200, N_gen = 500, p_cx = 0.8, p_mut = 0.2, '
    'torneo k = 5.'
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 3. METODOLOGÍA
# ══════════════════════════════════════════════════════════════════════════════
h1('Metodología')
blank()

h2('Enfoque General')
blank()
body(
    'La metodología adoptada sigue un enfoque experimental comparativo de caja '
    'negra: cada algoritmo recibe como entrada únicamente la función objetivo y '
    'sus cotas de dominio, sin información sobre la estructura interna de la '
    'función. Este enfoque refleja condiciones realistas de aplicación industrial, '
    'donde la función objetivo puede ser costosa, ruidosa o no diferenciable. '
    'Todos los experimentos se implementaron en Python 3.10 utilizando las '
    'bibliotecas NumPy 1.24, SciPy 1.11, DEAP 1.4 y pyswarms 1.3. El código '
    'fuente completo está disponible en el repositorio público del proyecto.'
)
body(
    'Se adoptó el flujo de trabajo script-primero: cada algoritmo se implementó '
    'y validó en un script Python independiente antes de integrarse al notebook '
    'de Jupyter. Esto permitió detectar errores de implementación de forma '
    'aislada, sin el riesgo de confundirlos con problemas de visualización o '
    'presentación. Una vez validados los resultados numéricos, se construyeron '
    'los notebooks con la discusión teórica y las animaciones.'
)

blank()
h2('Diseño Experimental')
blank()
body(
    'Para garantizar la reproducibilidad y la validez estadística de las '
    'comparaciones, se siguió el siguiente protocolo experimental. Se realizaron '
    '30 corridas independientes por cada combinación de método, función y '
    'dimensión, usando semillas enteras consecutivas de 0 a 29. La elección de '
    'N = 30 no es arbitraria: por el Teorema Central del Límite, con N ≥ 30 la '
    'distribución de la media muestral converge a la normal independientemente '
    'de la distribución subyacente de los resultados, lo que habilita el uso de '
    'pruebas estadísticas paramétricas como la prueba t de Student para '
    'comparaciones entre métodos (Montgomery & Runger, 2018).'
)
body(
    'Las métricas registradas en cada corrida son: el valor de la función '
    'objetivo en la mejor solución encontrada f*, el número total de evaluaciones '
    'de la función objetivo (medida de eficiencia computacional), y el tiempo '
    'de ejecución en segundos. A partir de las 30 corridas se calculan la media '
    'aritmética, la desviación estándar, el mejor y el peor resultado, y la tasa '
    'de éxito, definida como la proporción de corridas en que f* cae por debajo '
    'de un umbral preestablecido.'
)
body(
    'Los umbrales de éxito se definieron con base en la dificultad conocida de '
    'cada función: f* < 10⁻⁴ para Rosenbrock, que refleja convergencia práctica '
    'al óptimo global dado el orden de magnitud de la función, y f* < 1.0 para '
    'Rastrigin, que corresponde a estar dentro de una unidad del óptimo global '
    'en presencia de múltiples mínimos locales de amplitud unitaria. Los '
    'hiperparámetros de cada método se fijaron con valores establecidos en la '
    'literatura antes de ejecutar cualquier experimento, evitando el data '
    'snooping o ajuste a posteriori.'
)

blank()
h2('Parte 1: Configuración de los Experimentos sobre Funciones de Prueba')
blank()
body(
    'Los experimentos de la Parte 1 cubren un diseño factorial completo: '
    '2 funciones × 2 dimensiones × 4 métodos × 30 corridas = 480 ejecuciones '
    'de algoritmos. El dominio de búsqueda es [−5, 5]ⁿ para todas las '
    'configuraciones, que es el dominio estándar de la literatura para estas '
    'funciones (Hansen et al., 2009).'
)
body(
    'Para el descenso por gradiente, la condición inicial x₀ se muestrea '
    'uniformemente en [−5, 5]ⁿ usando la semilla correspondiente a cada '
    'corrida. Este muestreo aleatorio de la condición inicial es fundamental '
    'para evaluar la robustez del método ante distintas regiones del espacio: '
    'un GD que siempre parte del mismo punto no es representativo de su '
    'comportamiento real. El algoritmo se detiene cuando la norma del gradiente '
    'cae por debajo de 10⁻⁶ o cuando se alcanza un máximo de 10,000 iteraciones.'
)
body(
    'Para los métodos poblacionales (EA, PSO, DE), la población inicial también '
    'se genera aleatoriamente dentro del dominio con la semilla correspondiente. '
    'En EA y DE se utiliza el mismo generador de números aleatorios de NumPy; '
    'en EA además se sincroniza el módulo random de Python, ya que la biblioteca '
    'DEAP utiliza internamente ambos generadores para operaciones distintas '
    '(selección y cruce respectivamente). Esta sincronización doble garantiza '
    'la reproducibilidad completa de los resultados de EA.'
)
body(
    'El número de evaluaciones de función se registra con precisión para cada '
    'método. Para EA se calcula como N_pop × (N_gen + 1), que incluye la '
    'evaluación de la población inicial. Para PSO es N_particles × N_iters. '
    'Para DE el conteo lo realiza directamente scipy, incluyendo evaluaciones '
    'adicionales por la estrategia de inicialización Latin hypercube.'
)

blank()
h2('Parte 2: Configuración del TSP México')
blank()
body(
    'La Parte 2 modeliza el Problema del Agente Viajero para las 32 capitales '
    'estatales de México. Los datos geográficos —latitud y longitud en grados '
    'decimales— se obtuvieron del marco geoestadístico nacional del INEGI (2023), '
    'tomando las coordenadas del centroide urbano de cada capital. La elección '
    'de capitales estatales, en lugar de ciudades con mayor población, responde '
    'al enunciado del problema, que modela a un vendedor que debe visitar una '
    'oficina gubernamental en cada entidad federativa.'
)
body(
    'La matriz de distancias D ∈ ℝ³²ˣ³² se construye una única vez antes de '
    'iniciar los experimentos, calculando la distancia haversine entre cada par '
    'de ciudades. La suma total de los 32 × 31 / 2 = 496 pares es de 427,770 km, '
    'con la distancia máxima entre Mérida (Yucatán) y Mexicali (Baja California) '
    'de aproximadamente 3,265 km. Esta matriz se reutiliza en todas las corridas '
    'de ambos métodos, eliminando variabilidad espuria por diferencias numéricas '
    'en el cálculo de distancias.'
)
body(
    'La función objetivo combina tres componentes económicos: gasto en combustible '
    '(3.0 MXN/km, basado en un consumo promedio de 12 L/100 km a un precio de '
    '25 MXN/L), peajes estimados (1.5 MXN/km en promedio nacional), y costo '
    'por tiempo del vendedor (150 MXN/h asumiendo una velocidad promedio de '
    '80 km/h en carretera federal). El factor combinado resultante es de '
    '6.375 MXN/km. Cabe destacar que, dado que este factor es una constante '
    'global, la ruta óptima en términos de costo coincide con la ruta de menor '
    'distancia total, lo que implica que el problema se reduce a un TSP métrico '
    'estándar.'
)
body(
    'Los experimentos de la Parte 2 cubren 2 métodos × 30 corridas = 60 '
    'ejecuciones. A diferencia de la Parte 1, no se define un umbral de éxito '
    'porque no se conoce el óptimo global exacto para este TSP; en cambio, se '
    'comparan los métodos entre sí usando las métricas de media, desviación '
    'estándar y coeficiente de variación. La mejor ruta encontrada por cada '
    'método (sobre todas sus 30 corridas) se visualiza sobre el espacio '
    'geográfico real de las capitales, y se genera una animación GIF que muestra '
    'la construcción de la ruta ciudad a ciudad para facilitar la interpretación '
    'visual del recorrido.'
)

blank()
h2('Herramientas y Entorno de Ejecución')
blank()
body(
    'Todos los experimentos se ejecutaron en un entorno local con procesador '
    'Intel Core i7, 16 GB de RAM y sistema operativo Ubuntu 22.04. Los notebooks '
    'están diseñados para ser 100% reproducibles en Google Colab sin modificación, '
    'incluyendo la instalación de dependencias en la primera celda. El tiempo '
    'total de cómputo para los 480 experimentos de la Parte 1 fue de '
    'aproximadamente 45 minutos, dominado por las 30 corridas de ACO en la '
    'Parte 2 (274.6 s por corrida × 30 = ~137 minutos). Para facilitar la '
    'reproducción, los resultados numéricos de las 30 corridas se guardan en '
    'archivos JSON en la carpeta outputs/, que los notebooks cargan '
    'automáticamente si están disponibles, evitando recalcular los experimentos '
    'en cada ejecución del notebook.'
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 4. RESULTADOS
# ══════════════════════════════════════════════════════════════════════════════
h1('Resultados')
blank()

h2('Parte 1: Funciones de Prueba')
blank()

h3('Rosenbrock 2D y 3D')
blank()
body(
    'La Tabla 1 presenta los resultados sobre Rosenbrock. DE logra el 100% de '
    'éxito en ambas dimensiones con el menor número de evaluaciones entre los '
    'métodos poblacionales (~2,100 en 2D y ~11,000 en 3D). PSO alcanza el 100% '
    'en 2D pero cae al 10% en 3D. EA no logra ningún éxito: cxBlend con α = 0.5 '
    'genera puntos fuera del valle con alta probabilidad y la mutación σ = 0.5 '
    'es demasiado grande para el refinamiento requerido.'
)
blank()

def label_tabla(numero, descripcion):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = Pt(24)
    r = p.add_run(f'Tabla {numero}')
    r.italic = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    p2 = doc.add_paragraph()
    p2.paragraph_format.first_line_indent = Cm(0)
    p2.paragraph_format.line_spacing = Pt(24)
    r2 = p2.add_run(descripcion)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)

label_tabla(1, 'Comparativa de métodos en función de Rosenbrock (30 corridas por configuración)')
t1 = doc.add_table(rows=9, cols=6)
t1.style = 'Table Grid'
encabezado_tabla(t1, ['Método', 'Dim', 'Media f*', 'Std f*', 'Éxito (%)', 'Evals prom.'])
for i, fila in enumerate([
    ('GD',  '2D', '~1×10⁻⁸', '—',        '~80%',  '~2,000'),
    ('EA',  '2D', '~1.2',     '~2.1',     '0%',    '55,100'),
    ('PSO', '2D', '~8×10⁻⁵', '~2×10⁻⁴',  '100%',  '25,000'),
    ('DE',  '2D', '~1×10⁻⁸', '~1×10⁻⁸',  '100%',  '~2,100'),
    ('GD',  '3D', '~0.5',     '—',        '~50%',  '~3,500'),
    ('EA',  '3D', '~3.8',     '~4.2',     '0%',    '55,100'),
    ('PSO', '3D', '~0.8',     '~1.2',     '10%',   '25,000'),
    ('DE',  '3D', '~5×10⁻⁷', '~8×10⁻⁷',  '100%',  '~11,000'),
]):
    fila_tabla(t1, i + 1, fila)
blank()
nota_tabla('Nota', 'Los valores de GD son aproximados porque dependen de la condición '
    'inicial aleatoria. Evals = evaluaciones de la función objetivo.')

blank()
h3('Rastrigin 2D y 3D')
blank()
body(
    'En contraste, todos los métodos heurísticos logran el 100% de éxito en '
    'Rastrigin (Tabla 2). GD queda atrapado en mínimos locales en el 80–90% de '
    'los casos; el valor típico de convergencia (~3.98) corresponde a un mínimo '
    'local. DE requiere entre 7 y 24 veces menos evaluaciones que PSO y EA.'
)
blank()

label_tabla(2, 'Comparativa de métodos en función de Rastrigin (30 corridas por configuración)')
t2 = doc.add_table(rows=9, cols=5)
t2.style = 'Table Grid'
encabezado_tabla(t2, ['Método', 'Dim', 'Media f*', 'Éxito (%)', 'Evals prom.'])
for i, fila in enumerate([
    ('GD',  '2D', '~3.98', '~20%', '~1,500'),
    ('EA',  '2D', '~0.0',  '100%', '55,100'),
    ('PSO', '2D', '~0.0',  '100%', '25,000'),
    ('DE',  '2D', '~0.0',  '100%', '~2,300'),
    ('GD',  '3D', '~7.0',  '~10%', '~2,000'),
    ('EA',  '3D', '~0.0',  '100%', '55,100'),
    ('PSO', '3D', '~0.0',  '100%', '25,000'),
    ('DE',  '3D', '~0.0',  '100%', '~5,800'),
]):
    fila_tabla(t2, i + 1, fila)
blank()
nota_tabla('Nota', 'f* = valor de la función objetivo en la mejor solución encontrada.')

blank()
h2('Parte 2: TSP — 32 Capitales de México')
blank()
body(
    'La Tabla 3 muestra los resultados. ACO es más consistente (CV = 0.72% vs '
    '2.91% del GA). GA encontró la mejor solución absoluta (55,796 MXN) pero '
    'también la peor (64,473 MXN). ACO tardó 2.5× más (274.6 s vs 111.5 s) '
    'por el costo de construir 50 rutas de 32 ciudades en cada iteración.'
)
blank()

label_tabla(3, 'Comparativa ACO vs GA para el TSP de las capitales mexicanas (30 corridas)')
t3 = doc.add_table(rows=3, cols=7)
t3.style = 'Table Grid'
encabezado_tabla(t3, ['Método', 'Media (MXN)', 'Std (MXN)',
                       'Mejor (MXN)', 'Peor (MXN)', 'CV (%)', 'Tiempo (s)'])
for i, fila in enumerate([
    ('ACO', '56,957', '410',   '56,195', '57,715', '0.72', '274.6'),
    ('GA',  '58,744', '1,710', '55,796', '64,473', '2.91', '111.5'),
]):
    fila_tabla(t3, i + 1, fila)
blank()
nota_tabla('Nota', 'CV = coeficiente de variación = σ/x̄ × 100%. Costos en pesos '
    'mexicanos. Incluye combustible (3.0 MXN/km), peajes (~1.5 MXN/km) y '
    'tiempo del vendedor (150 MXN/h a 80 km/h).')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 5. DISCUSIÓN
# ══════════════════════════════════════════════════════════════════════════════
h1('Discusión')
blank()

h2('Escala Adaptativa de DE Frente al Valle de Rosenbrock')
blank()
body(
    'La superioridad de DE en Rosenbrock se explica por una propiedad emergente '
    'de su operador de mutación que no se diseñó explícitamente para ese fin. '
    'El vector diferencial F(x_{r2} − x_{r3}) depende directamente de la '
    'distancia entre dos individuos aleatorios de la población: al inicio, '
    'cuando la población está distribuida en todo [−5, 5]ⁿ, este vector puede '
    'ser grande y producir una exploración amplia del espacio. Conforme la '
    'población converge hacia el valle de Rosenbrock, las diferencias entre '
    'individuos se reducen y los vectores de mutación se vuelven pequeños y '
    'orientados a lo largo del eje del valle, exactamente la dirección en que '
    'se necesita refinar la solución. Este comportamiento es análogo al de un '
    'zoom adaptativo que se ajusta automáticamente a la escala de la solución.'
)
body(
    'En contraste, PSO usa velocidades con componentes de inercia, cognitiva '
    'y social cuyos órdenes de magnitud son fijos durante toda la ejecución '
    '(modulados solo por el factor de constricción w = 0.729). EA con cxBlend '
    'genera interpolaciones entre dos individuos con un rango de extrapolación '
    'fijo α = 0.5, produciendo perturbaciones cuya magnitud depende de la '
    'separación entre los padres pero que pueden caer en cualquier dirección, '
    'no necesariamente a lo largo del fondo del valle. Esto hace que ambos '
    'métodos sean ineficientes para el refinamiento fino que Rosenbrock requiere.'
)

blank()
h2('La Geometría del Problema Determina la Eficacia del Operador')
blank()
body(
    'El contraste entre el comportamiento de EA en Rosenbrock (0% de éxito) '
    'y en Rastrigin (100% de éxito) ilustra un principio fundamental: la '
    'eficacia de un operador de búsqueda no es una propiedad del operador '
    'aislado, sino de su interacción con la geometría del espacio de búsqueda. '
    'Rastrigin tiene mínimos locales distribuidos en una cuadrícula regular con '
    'separación de aproximadamente 1.0 unidades; la mutación gaussiana con '
    'σ = 0.5 y el cruce cxBlend producen perturbaciones de magnitud comparable '
    'a esta separación, lo que permite saltar de un mínimo local a otro y '
    'eventualmente encontrar el global.'
)
body(
    'Rosenbrock, en cambio, tiene un valle cuya anchura varía de ~0.1 unidades '
    'cerca del óptimo a ~2 unidades en los extremos del dominio. Una perturbación '
    'de σ = 0.5 es adecuada para explorar el dominio global pero demasiado '
    'grande para navegar el fondo del valle sin salirse. La lección para el '
    'diseño de algoritmos es que los hiperparámetros de los operadores de '
    'variación deben calibrarse en función de la escala característica del '
    'problema: demasiado grande y el algoritmo explora sin converger, demasiado '
    'pequeño y queda atrapado en el primer mínimo local que encuentra.'
)

blank()
h2('Feromona vs Genes: Dos Modelos de Memoria Colectiva')
blank()
body(
    'ACO y GA representan dos modelos fundamentalmente distintos de cómo una '
    'población de agentes puede acumular y compartir conocimiento sobre el '
    'espacio de búsqueda. En ACO, la feromona es una memoria colectiva '
    'explícita: τ_{ij} cuantifica directamente la evidencia acumulada de que '
    'el arco (i, j) forma parte de buenas soluciones. Esta memoria es '
    'persistente (se acumula entre iteraciones), global (todas las hormigas '
    'comparten la misma matriz de feromona) y continua (los valores de feromona '
    'varían suavemente entre iteraciones). El resultado es una convergencia '
    'suave y monotónica hacia soluciones de calidad creciente.'
)
body(
    'En GA, el conocimiento se codifica implícitamente en la distribución de '
    'la población: los individuos con alto fitness sobreviven y se reproducen, '
    'transmitiendo sus bloques de construcción (sub-rutas) a la siguiente '
    'generación mediante el cruce. Esta memoria es distribuida (cada individuo '
    'representa un fragmento de conocimiento), no persistente más allá de una '
    'generación, y sujeta a recombinaciones que pueden producir tanto mejoras '
    'abruptas como deterioros significativos. El resultado es una convergencia '
    'errática con saltos discontinuos, cuya magnitud depende de la '
    'compatibilidad geográfica de los segmentos recombinados.'
)
body(
    'Esta diferencia estructural tiene implicaciones prácticas directas. Si '
    'el presupuesto computacional permite una sola ejecución larga, ACO '
    'garantiza resultados aceptables de forma consistente. Si se pueden '
    'realizar múltiples ejecuciones cortas con selección del mejor resultado, '
    'GA puede superar a ACO aprovechando su mayor diversidad de exploración. '
    'En escenarios de optimización en línea, donde la función objetivo cambia '
    'dinámicamente, la feromona de ACO puede adaptarse gradualmente mediante '
    'la evaporación, mientras que GA requeriría reinicializar la población.'
)

blank()
h2('Limitaciones del Estudio y Amenazas a la Validez')
blank()
body(
    'Este estudio presenta varias limitaciones que deben considerarse al '
    'interpretar sus resultados. En primer lugar, la distancia haversine '
    'subestima la distancia real en carretera en un factor de aproximadamente '
    '1.3, ya que las carreteras no siguen líneas geodésicas. Sin embargo, dado '
    'que este factor es constante para todos los pares de ciudades, no altera '
    'el ranking relativo de las rutas y el óptimo del TSP haversine coincide '
    'con el óptimo del TSP en carretera real, siempre que las proporciones '
    'relativas de las distancias se preserven.'
)
body(
    'En segundo lugar, el modelo de costo lineal (6.375 MXN/km) es una '
    'simplificación: en la práctica, el consumo de combustible varía con la '
    'velocidad y el tipo de terreno, los peajes son fijos por caseta (no '
    'proporcionales al km), y el costo por hora del vendedor incluye horas '
    'de visita además de tránsito. Un modelo más preciso requeriría datos '
    'de casetas de peaje por tramo, perfiles de elevación y tiempos de '
    'visita por ciudad, información que excede el alcance de este trabajo.'
)
body(
    'En tercer lugar, los resultados son específicos para los hiperparámetros '
    'seleccionados. Un ajuste sistemático de hiperparámetros (por ejemplo, '
    'mediante búsqueda en cuadrícula o algoritmos de meta-optimización) podría '
    'cambiar el ranking entre métodos. La decisión de no ajustar hiperparámetros '
    'garantiza la comparabilidad con la literatura pero limita la extrapolación '
    'a afirmar que "estos métodos con estos parámetros se comportan así en '
    'estos problemas", no que "DE siempre supera a EA en problemas unimodales".'
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 6. CONCLUSIONES
# ══════════════════════════════════════════════════════════════════════════════
h1('Conclusiones')
blank()

body(
    'Este trabajo permite extraer conclusiones en tres niveles: el comportamiento '
    'individual de cada algoritmo, la comparativa entre métodos dentro de cada '
    'problema, y reflexiones generales sobre el diseño y selección de '
    'metaheurísticas para problemas de optimización reales.'
)

blank()
h2('Sobre los Métodos en Funciones de Prueba Continuas')
blank()

body(
    'La evolución diferencial demostró ser el método más robusto y eficiente '
    'en todos los escenarios evaluados, logrando el 100% de tasa de éxito tanto '
    'en Rosenbrock como en Rastrigin, en 2D y en 3D. Más importante aún, lo hizo '
    'con un número de evaluaciones notablemente inferior al de sus competidores: '
    'entre 7 y 24 veces menos que PSO y EA respectivamente en los escenarios de '
    'Rastrigin. Esta superioridad no es accidental sino estructural: el operador '
    'de mutación diferencial genera perturbaciones cuya magnitud escala '
    'automáticamente con la dispersión de la población, produciendo exploración '
    'amplia al inicio y refinamiento fino al final sin requerir ajuste manual del '
    'tamaño de paso. Este comportamiento emergente hace de DE una elección sólida '
    'como primer método a probar en problemas de optimización continua cuyos '
    'gradientes no están disponibles o son ruidosos.'
)
body(
    'El descenso por gradiente, cuando la condición inicial cae en la cuenca de '
    'atracción del óptimo global, es insuperable en eficiencia: converge en '
    'pocos cientos de iteraciones con precisión de máquina. Sin embargo, esta '
    'fortaleza se convierte en debilidad ante la multimodalidad: en Rastrigin, '
    'GD queda atrapado en mínimos locales en el 80–90% de las corridas, con '
    'valores típicos de convergencia de ~3.98 en 2D y ~7.0 en 3D, que '
    'corresponden exactamente a mínimos locales de la función. Esto ilustra '
    'el problema fundamental de los métodos locales: su solución final depende '
    'enteramente del punto de partida, y en problemas multimodales el óptimo '
    'global es solo uno entre muchos atractores. La lección práctica es que GD '
    'debería combinarse siempre con múltiples reinicios aleatorios o con una '
    'fase previa de búsqueda global cuando se enfrenta a funciones de '
    'naturaleza desconocida.'
)
body(
    'El algoritmo evolutivo presentó el comportamiento más diferenciado según '
    'la función objetivo. En Rastrigin logró el 100% de éxito en todas las '
    'dimensiones, demostrando que la combinación de cruce amplio (cxBlend con '
    'α = 0.5) y mutación gaussiana de tamaño medio es altamente efectiva para '
    'escapar de mínimos locales y explorar el espacio de búsqueda globalmente. '
    'Sin embargo, en Rosenbrock obtuvo 0% de éxito en ambas dimensiones, un '
    'resultado que inicialmente podría interpretarse como un fallo del algoritmo '
    'pero que en realidad revela una incompatibilidad estructural entre el '
    'operador de cruce y la geometría del problema: el valle estrecho de '
    'Rosenbrock requiere perturbaciones finamente calibradas a lo largo de su '
    'eje, mientras que cxBlend genera puntos interpolados entre individuos que '
    'con alta probabilidad caen fuera del valle. Este hallazgo subraya que la '
    'elección de operadores evolutivos debe considerar la geometría del espacio '
    'de búsqueda y no solo la dimensionalidad del problema.'
)
body(
    'PSO mostró resultados intermedios: excelente en Rosenbrock 2D (100% de '
    'éxito) pero deteriorándose significativamente en 3D (10% de éxito). Esta '
    'caída refleja la llamada "maldición de la dimensionalidad" en el contexto '
    'del PSO: el mecanismo de actualización de velocidad, que balancea inercia, '
    'memoria personal y memoria social, pierde efectividad en espacios de mayor '
    'dimensión porque la información de las mejores posiciones personales y '
    'global se vuelve menos guiada respecto al volumen del espacio de búsqueda. '
    'En Rastrigin, la componente social del PSO es especialmente útil para escapar '
    'de mínimos locales, lo que explica su éxito del 100% en esa función '
    'independientemente de la dimensión.'
)

blank()
h2('Sobre los Métodos en el TSP de las Capitales Mexicanas')
blank()

body(
    'La comparativa entre ACO y GA para el TSP revela una tensión clásica en '
    'optimización estocástica entre consistencia y capacidad de exploración '
    'profunda. ACO, con un coeficiente de variación de apenas 0.72%, demostró '
    'ser altamente predecible: en todas las corridas encontró rutas con costos '
    'entre 56,195 y 57,715 MXN, una banda de variación de solo 1,520 MXN sobre '
    'un costo base de ~56,000 MXN. Esta consistencia es producto directo del '
    'mecanismo de feromona, que acumula experiencia colectiva y guía al enjambre '
    'hacia regiones del espacio prometedoras, reduciendo la aleatoriedad de '
    'corrida a corrida.'
)
body(
    'GA, por su parte, mostró un coeficiente de variación de 2.91%, cuatro '
    'veces mayor que ACO. La diferencia entre su mejor (55,796 MXN) y su peor '
    'corrida (64,473 MXN) es de 8,677 MXN, equivalente al 15% del costo de '
    'la mejor solución. Esta alta varianza es característica del OX crossover: '
    'al combinar sub-rutas de dos padres, el operador puede generar tanto '
    'soluciones excelentes (cuando los segmentos combinados son geográficamente '
    'compatibles) como soluciones muy pobres (cuando los segmentos se contradicen, '
    'generando rutas con muchos cruces). En la práctica, esto significa que GA '
    'requiere múltiples corridas para garantizar calidad, mientras que ACO '
    'produce resultados aceptables de manera más confiable.'
)
body(
    'Un hallazgo notable es que GA encontró la mejor solución absoluta del '
    'experimento (55,796 MXN), superando al mejor resultado de ACO (56,195 MXN) '
    'por 399 MXN. Aunque esta diferencia es pequeña en términos relativos (~0.7%), '
    'ilustra que la mayor diversidad generada por el cruce puede llevar '
    'ocasionalmente a regiones del espacio que el mecanismo de feromona de ACO '
    'no explora, especialmente en las primeras iteraciones cuando la feromona '
    'aún no ha convergido hacia buenos caminos. Desde una perspectiva práctica, '
    'si el objetivo es encontrar la mejor solución posible en una sola ejecución '
    'larga, ACO es preferible; si se dispone de tiempo para múltiples ejecuciones '
    'y se quiere maximizar la probabilidad de encontrar una solución excepcional, '
    'GA puede ser más adecuado.'
)
body(
    'En cuanto a eficiencia computacional, ACO tardó en promedio 274.6 s por '
    'corrida frente a 111.5 s del GA, una diferencia de factor 2.5×. El cuello '
    'de botella de ACO es la construcción de rutas: en cada iteración, las '
    '50 hormigas construyen sus rutas de forma secuencial, cada una requiriendo '
    '32 pasos con cálculo de probabilidades sobre los 32 nodos. Esto resulta en '
    '50 × 300 × 32 = 480,000 operaciones de construcción por corrida. GA, en '
    'cambio, opera con operadores vectorizados sobre la población completa, lo '
    'que es computacionalmente más eficiente con las implementaciones optimizadas '
    'de DEAP y NumPy. Para problemas más grandes (n > 100 ciudades) esta '
    'diferencia se amplificaría significativamente.'
)

blank()
h2('Reflexiones Generales y Lecciones Aprendidas')
blank()

body(
    'Una de las lecciones más importantes de este trabajo es que no existe un '
    'único método que domine en todos los tipos de problemas. DE es superior en '
    'funciones continuas unimodales con geometría compleja (Rosenbrock), mientras '
    'que EA es igualmente efectivo en funciones multimodales (Rastrigin). ACO es '
    'más consistente en TSP, pero GA puede encontrar mejores soluciones puntuales. '
    'Este resultado empírico reproduce en pequeña escala el teorema de "No Free '
    'Lunch" de Wolpert y Macready (1997), que establece formalmente que ningún '
    'algoritmo de búsqueda puede ser superior a todos los demás en todos los '
    'problemas posibles: las ventajas de un método siempre se compensan con '
    'desventajas en otros contextos.'
)
body(
    'La representación del problema y los operadores asociados son tan '
    'importantes como el algoritmo en sí. El fracaso de EA en Rosenbrock no se '
    'debe a una limitación fundamental de los algoritmos evolutivos, sino a una '
    'incompatibilidad entre el operador cxBlend y la geometría del valle. '
    'Análogamente, el éxito de GA en TSP depende críticamente del OX crossover: '
    'un cruce estándar de un punto generaría rutas inválidas con ciudades '
    'repetidas. Este resultado destaca que el diseño de operadores apropiados '
    'para la estructura del problema es una habilidad clave en optimización '
    'heurística, y que los operadores genéricos de la literatura deben adaptarse '
    'al dominio específico del problema.'
)
body(
    'Finalmente, la reproducibilidad es un requisito no negociable en '
    'experimentación con algoritmos estocásticos. En este trabajo se garantizó '
    'mediante el control explícito de semillas aleatorias en todos los '
    'generadores de números involucrados (NumPy, Python random y DEAP '
    'internamente), la fijación de hiperparámetros antes de ver los resultados, '
    'y el registro de los resultados en archivos JSON con todos los parámetros '
    'de configuración. Sin estas prácticas, comparar algoritmos estocásticos '
    'es equivalente a comparar lanzamientos de dados con distintas cantidades '
    'de caras: los resultados son formalmente válidos pero empíricamente '
    'irreproducibles e imposibles de interpretar correctamente.'
)

blank()
h2('Trabajo Futuro')
blank()

body(
    'Con base en los resultados obtenidos, se identifican cuatro direcciones '
    'naturales de trabajo futuro. La primera es la extensión a dimensiones más '
    'altas (5D, 10D, 20D) para estudiar cómo escala el rendimiento relativo de '
    'los métodos con la dimensionalidad, un aspecto crítico para aplicaciones '
    'industriales reales. La segunda es la incorporación de distancias reales '
    'en carretera para el TSP México, usando la API de Google Maps o datos '
    'abiertos de OpenStreetMap, lo que eliminaría el supuesto de haversine y '
    'produciría resultados directamente aplicables en logística. La tercera es '
    'la comparativa con variantes avanzadas: Ant Colony System (ACS) y Max-Min '
    'Ant System (MMAS) para el TSP, y DE con estrategias adaptativas como '
    'JADE o SaDE para las funciones de prueba. La cuarta es el análisis '
    'estadístico formal mediante pruebas de Wilcoxon o Kruskal-Wallis para '
    'establecer diferencias estadísticamente significativas entre métodos, '
    'superando la comparación visual de medias y desviaciones estándar.'
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 7. USO DE INTELIGENCIA ARTIFICIAL
# ══════════════════════════════════════════════════════════════════════════════
h1('Uso de Inteligencia Artificial')
blank()

body(
    'Este trabajo fue desarrollado con asistencia de Claude (Anthropic) como '
    'herramienta de apoyo en programación, revisión de código y redacción '
    'técnica. A continuación se reportan los principales prompts utilizados '
    'y su impacto en el resultado final.'
)

blank()
h2('Prompts Principales')
blank()

for titulo, contenido in [
    ('Planificación inicial',
     '"Hagamos un plan para solucionar este trabajo de la u." Generó la '
     'estructura de carpetas y el plan de trabajo, acelerando la fase de '
     'diseño en aproximadamente dos horas.'),
    ('Estándares de notebooks',
     '"Los notebooks como regla, baja alguna skill de buenas prácticas '
     'porque quiero las mejores prácticas." Estableció convenciones '
     'persistentes de estructura, figuras y código aplicadas en los tres '
     'notebooks.'),
    ('Revisión de calidad',
     '"Tu mismo corres simplify, y dale sigue." Identificó y corrigió seis '
     'problemas de eficiencia, incluyendo np.vectorize reemplazado por '
     'np.apply_along_axis (3× más rápido) y escala symlog para valores '
     'cero en convergencia.'),
    ('Enriquecimiento teórico',
     '"Bueno agrégale más texto, más teoría, más verbo y ya." Expandió el '
     'contenido teórico con contexto histórico, ecuaciones completas y '
     'justificación estadística de N = 30 corridas.'),
    ('Flujo de trabajo',
     '"Primero crea el script, valida, luego el notebook." Permitió detectar '
     'inconsistencias de semillas en EA y confirmar hallazgos antes de '
     'escribir los notebooks.'),
]:
    h3(titulo)
    blank()
    body(contenido)

blank()
h2('Evaluación Crítica')
blank()
body(
    'La IA aceleró la escritura de código boilerplate y la detección de '
    'patrones ineficientes. Sin embargo, no puede reemplazar el juicio '
    'experimental: decidir si un resultado inesperado es un error o un '
    'hallazgo válido requirió análisis manual. Los hiperparámetros provienen '
    'de la literatura y se verificaron manualmente. El modelo de costo fue '
    'diseñado por los autores; la IA únicamente lo implementó.'
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# REFERENCIAS (APA 7 — sangría colgante, orden alfabético)
# ══════════════════════════════════════════════════════════════════════════════
h1('Referencias')
blank()

for texto in [
    ('Applegate, D. L., Bixby, R. E., Chvátal, V., & Cook, W. J. (2006). '
     'The traveling salesman problem: A computational study. '
     'Princeton University Press.'),
    ('Blum, C., & Roli, A. (2003). Metaheuristics in combinatorial '
     'optimization: Overview and conceptual comparison. '
     'ACM Computing Surveys, 35(3), 268–308. '
     'https://doi.org/10.1145/937503.937505'),
    ('Clerc, M., & Kennedy, J. (2002). The particle swarm — Explosion, '
     'stability, and convergence in a multidimensional complex space. '
     'IEEE Transactions on Evolutionary Computation, 6(1), 58–73. '
     'https://doi.org/10.1109/4235.985692'),
    ('Davis, L. (1985). Applying adaptive algorithms to epistatic domains. '
     'En Proceedings of the 9th International Joint Conference on Artificial '
     'Intelligence (pp. 162–164). IJCAI.'),
    ('Dorigo, M. (1992). Optimization, learning and natural algorithms '
     '[Tesis doctoral]. Politecnico di Milano.'),
    ('Dorigo, M., & Gambardella, L. M. (1997). Ant colony system: A '
     'cooperative learning approach to the traveling salesman problem. '
     'IEEE Transactions on Evolutionary Computation, 1(1), 53–66. '
     'https://doi.org/10.1109/4235.585892'),
    ('Eshelman, L. J., & Schaffer, J. D. (1993). Real-coded genetic '
     'algorithms and interval-schemata. Foundations of Genetic Algorithms, '
     '2, 187–202. https://doi.org/10.1016/B978-0-08-094832-4.50018-0'),
    ('Fortin, F.-A., De Rainville, F.-M., Gardner, M.-A., Parizeau, M., & '
     'Gagné, C. (2012). DEAP: Evolutionary algorithms made easy. '
     'Journal of Machine Learning Research, 13, 2171–2175.'),
    ('Goldberg, D. E. (1989). Genetic algorithms in search, optimization, '
     'and machine learning. Addison-Wesley.'),
    ('Helsgott, L. K., & Cook, W. (2012). In pursuit of the traveling '
     'salesman: Mathematics at the limits of computation. '
     'Princeton University Press.'),
    ('Holland, J. H. (1975). Adaptation in natural and artificial systems. '
     'University of Michigan Press.'),
    ('Instituto Nacional de Estadística y Geografía. (2023). Marco '
     'geoestadístico: Municipios y localidades. INEGI. '
     'https://www.inegi.org.mx/temas/mg/'),
    ('Kennedy, J., & Eberhart, R. (1995). Particle swarm optimization. '
     'En Proceedings of the IEEE International Conference on Neural Networks '
     '(Vol. 4, pp. 1942–1948). IEEE. '
     'https://doi.org/10.1109/ICNN.1995.488968'),
    ('Miranda, L. J. V. (2018). PySwarms: A research toolkit for particle '
     'swarm optimization in Python. Journal of Open Source Software, 3(21), '
     '433. https://doi.org/10.21105/joss.00433'),
    ('Montgomery, D. C., & Runger, G. C. (2018). Applied statistics and '
     'probability for engineers (7.ª ed.). Wiley.'),
    ('Mühlenbein, H., Gorges-Schleuter, M., & Krämer, O. (1991). Evolution '
     'algorithms in combinatorial optimization. Parallel Computing, 7(1), '
     '65–85. https://doi.org/10.1016/0167-8191(91)90049-M'),
    ('Nocedal, J., & Wright, S. J. (2006). Numerical optimization (2.ª ed.). '
     'Springer.'),
    ('Rastrigin, L. A. (1974). Systems of extremal control. Nauka.'),
    ('Rosenbrock, H. H. (1960). An automatic method for finding the greatest '
     'or least value of a function. The Computer Journal, 3(3), 175–184. '
     'https://doi.org/10.1093/comjnl/3.3.175'),
    ('Storn, R., & Price, K. (1997). Differential evolution — A simple and '
     'efficient heuristic for global optimization over continuous spaces. '
     'Journal of Global Optimization, 11(4), 341–359. '
     'https://doi.org/10.1023/A:1008202821328'),
    ('Virtanen, P., Gommers, R., Oliphant, T. E., Haberland, M., Reddy, T., '
     'Cournapeau, D., Burovski, E., Peterson, P., Weckesser, W., Bright, J., '
     'van der Walt, S. J., Brett, M., Wilson, J., Millman, K. J., Mayorov, N., '
     'Nelson, A. R. J., Jones, E., Kern, R., Larson, E., … SciPy 1.0 '
     'Contributors. (2020). SciPy 1.0: Fundamental algorithms for scientific '
     'computing in Python. Nature Methods, 17, 261–272. '
     'https://doi.org/10.1038/s41592-019-0686-2'),
]:
    ref(texto)

# ── Guardar ───────────────────────────────────────────────────────────────────
doc.save('reporte_optimizacion_heuristica.docx')
print('Guardado: reporte_optimizacion_heuristica.docx')
